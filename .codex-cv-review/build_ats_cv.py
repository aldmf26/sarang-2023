from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\User\Downloads\CV Muhammad Rizki Hidayat - ATS.docx"

FONT = "Arial"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)


def set_font(run, size=10.5, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_keep_with_next(paragraph, value=True):
    pPr = paragraph._p.get_or_add_pPr()
    el = pPr.find(qn("w:keepNext"))
    if value and el is None:
        pPr.append(OxmlElement("w:keepNext"))
    elif not value and el is not None:
        pPr.remove(el)


def section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_font(r, size=11.5, bold=True)
    return p


def role(doc, company, location, period, position, details):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{position} | {company}")
    set_font(r, size=10.5, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{location} | {period}")
    set_font(r, size=9.5, italic=True, color=GRAY)

    for detail in details:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(detail)
        set_font(r, size=10)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.05

h1 = doc.styles["Heading 1"]
h1.font.name = FONT
h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
h1.font.size = Pt(11.5)
h1.font.bold = True
h1.font.color.rgb = BLACK
h1.paragraph_format.space_before = Pt(7)
h1.paragraph_format.space_after = Pt(2.5)
h1.paragraph_format.keep_with_next = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
r = p.add_run("MUHAMMAD RIZKI HIDAYAT")
set_font(r, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Banjarmasin, Indonesia | 0859-1066-77777 | rizkimuhammad@gmail.com")
set_font(r, size=9.5)

section_heading(doc, "Profil Profesional")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run(
    "Lulusan Teknologi Laboratorium Medik Universitas 'Aisyiyah Yogyakarta dengan IPK 3,12 "
    "dan pengalaman praktik kerja di rumah sakit serta puskesmas. Memiliki pengalaman pendukung "
    "di bidang operasional minimarket dan gudang. Terbiasa bekerja secara rapi, cekatan, teratur, "
    "mengikuti prosedur, melayani pelanggan, dan berkolaborasi dalam tim."
)
set_font(r, size=10)

section_heading(doc, "Pengalaman Kerja")
role(
    doc,
    "Minimarket OMI KJS Mart UM",
    "Banjarmasin",
    "Desember 2025 - Sekarang",
    "Karyawan Minimarket",
    [
        "Mendukung kegiatan operasional harian minimarket dan pelayanan pelanggan.",
        "Membantu penataan barang serta menjaga kerapian area kerja.",
    ],
)
role(
    doc,
    "Koperasi Jasa Syariah Cangkal Becari UM",
    "Banjarmasin",
    "September 2024 - Desember 2025",
    "Helper Gudang",
    [
        "Membantu penataan dan pemindahan barang di area gudang.",
        "Menjaga kerapian area kerja dan melaksanakan tugas sesuai arahan.",
    ],
)
role(
    doc,
    "RS PKU Muhammadiyah Surakarta",
    "Solo",
    "Februari 2020 - Maret 2020",
    "Magang",
    ["Mengikuti praktik kerja dan alur pelayanan laboratorium di bawah supervisi pembimbing."],
)
role(
    doc,
    "Puskesmas Tempel II Yogyakarta",
    "Yogyakarta",
    "Desember 2019 - Januari 2020",
    "Magang",
    ["Mengikuti praktik kerja dan kegiatan pelayanan laboratorium sesuai arahan pembimbing."],
)

section_heading(doc, "Pendidikan")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Teknologi Laboratorium Medik | Universitas 'Aisyiyah Yogyakarta")
set_font(r, size=10.5, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Yogyakarta | IPK 3,12 - Predikat Sangat Memuaskan")
set_font(r, size=9.5, color=GRAY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("SMK Unggulan Husada | Banjarmasin")
set_font(r, size=10.5, bold=True)

section_heading(doc, "Keterampilan")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run(
    "Pelayanan pelanggan | Operasional minimarket | Penataan barang | Dukungan operasional gudang | "
    "Dasar-dasar laboratorium medik | Komunikasi | Kerja sama tim | Kerapian dan ketelitian"
)
set_font(r, size=10)

section_heading(doc, "Pengalaman Organisasi")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("HIMATELMA | Universitas 'Aisyiyah Yogyakarta | 2016 - 2018")
set_font(r, size=10.5, bold=True)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.18)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Berpartisipasi dalam kepanitiaan seminar nasional kesehatan dan kegiatan bakti sosial.")
set_font(r, size=10)

doc.core_properties.title = "CV Muhammad Rizki Hidayat - ATS"
doc.core_properties.subject = "Curriculum Vitae"
doc.core_properties.author = "Muhammad Rizki Hidayat"
doc.core_properties.keywords = (
    "ATS, operasional, minimarket, gudang, pelayanan pelanggan, laboratorium medik, teamwork"
)
doc.core_properties.comments = ""

doc.save(OUT)
print(OUT)
